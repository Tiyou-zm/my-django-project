from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\Administrator\Desktop")
SRC = ROOT / "能源管理平台用户说明书.md"
OUT = ROOT / "能源管理平台用户说明书.docx"


def set_run_font(run, name="宋体", size=12, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def add_field(paragraph, field_code: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r1 = OxmlElement("w:r")
    r1.append(begin)
    r2 = OxmlElement("w:r")
    r2.append(instr)
    r3 = OxmlElement("w:r")
    r3.append(separate)
    r4 = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "右键更新目录"
    r4.append(text)
    r5 = OxmlElement("w:r")
    r5.append(end)

    paragraph._p.append(r1)
    paragraph._p.append(r2)
    paragraph._p.append(r3)
    paragraph._p.append(r4)
    paragraph._p.append(r5)


def add_page_number(paragraph):
    add_field(paragraph, "PAGE")


def add_inline_runs(paragraph, text, default_font="宋体", default_size=12):
    text = text.replace("`", "")
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, default_font, default_size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, default_font, default_size)


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string("000000")

    for style_name, font_name, size in [
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
        ("List Bullet", "宋体", 12),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True if "Heading" in style_name else False
        style.font.color.rgb = RGBColor.from_string("000000")


def add_cover(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    set_run_font(run, "黑体", 22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("软件用户说明书")
    set_run_font(run, "宋体", 16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(220)
    run = meta.add_run("文档类型：软著附件格式排版稿")
    set_run_font(run, "宋体", 12)

    doc.add_page_break()


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("目录")
    set_run_font(run, "黑体", 18, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


def format_heading(doc: Document, level: int, text: str):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(10 if level == 3 else 14)
    p.paragraph_format.space_after = Pt(4 if level == 3 else 6)
    run = p.add_run(text)
    set_run_font(run, "黑体", 16 if level == 1 else 14 if level == 2 else 12, bold=True)


def format_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.84)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text)


def format_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def format_term_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(4.2), Cm(10.8)]

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.width = widths[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.35
            run = p.add_run(value)
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(run, "黑体", 11, bold=True)
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr.append(shd)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_run_font(run, "宋体", 11)
    doc.add_paragraph()


def parse_rows(table_lines):
    rows = []
    for raw in table_lines:
        if re.match(r"^\|(?:\s*:?-+:?\s*\|)+\s*$", raw):
            continue
        rows.append([c.strip() for c in raw.strip().strip("|").split("|")])
    return rows


def build():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style_document(doc)
    add_update_fields_on_open(doc)

    title_line = next((line for line in lines if line.startswith("# ")), "# 用户说明书")
    add_cover(doc, title_line[2:].strip())
    add_toc(doc)

    in_table = []
    first_title_consumed = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# ") and not first_title_consumed:
            first_title_consumed = True
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            in_table.append(stripped)
            continue

        if in_table:
            rows = parse_rows(in_table)
            if rows:
                format_term_table(doc, rows)
            in_table = []

        if not stripped or stripped == "---":
            continue

        if stripped.startswith("## "):
            format_heading(doc, 1, stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            format_heading(doc, 2, stripped[4:].strip())
            continue

        if stripped.startswith("# "):
            format_heading(doc, 3, stripped[2:].strip())
            continue

        if stripped.startswith("- "):
            format_bullet(doc, stripped[2:].strip())
            continue

        format_body(doc, stripped)

    if in_table:
        rows = parse_rows(in_table)
        if rows:
            format_term_table(doc, rows)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer)

    doc.save(OUT)


if __name__ == "__main__":
    build()
